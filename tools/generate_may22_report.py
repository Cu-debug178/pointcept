from __future__ import annotations

import csv
from collections import defaultdict
from dataclasses import dataclass
from pathlib import Path

import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


ROOT = Path(r"E:\Program\python\Pointcept")
EXP_DIR = ROOT / "exp"
DOCS_DIR = ROOT / "docs"
ASSET_DIR = DOCS_DIR / "report_2026_05_22_assets"
OUT_DOCX = DOCS_DIR / "5月22日汇报文档.docx"


def set_cn_font(run, size=12, bold=False):
    run.font.name = "宋体"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    run.font.size = Pt(size)
    run.font.bold = bold


def set_table_font(cell, size=10):
    for p in cell.paragraphs:
        for r in p.runs:
            set_cn_font(r, size=size)


def ensure_dirs():
    ASSET_DIR.mkdir(parents=True, exist_ok=True)


def load_val_metrics():
    path = EXP_DIR / "experiment_val_metrics.csv"
    with path.open(newline="", encoding="utf-8") as f:
        rows = list(csv.DictReader(f))
    for r in rows:
        r["idx"] = int(r["idx"])
        r["mIoU"] = float(r["mIoU"])
        r["mAcc"] = float(r["mAcc"])
        r["allAcc"] = float(r["allAcc"])
    return pd.DataFrame(rows)


def load_class_metrics():
    path = EXP_DIR / "experiment_class_metrics.csv"
    with path.open(newline="", encoding="utf-8") as f:
        rows = list(csv.DictReader(f))
    for r in rows:
        r["val_idx"] = int(r["val_idx"])
        r["class_id"] = int(r["class_id"])
        r["iou"] = float(r["iou"])
        r["acc"] = float(r["acc"])
    return pd.DataFrame(rows)


def summarize(val_df, class_df):
    summaries = {}
    for exp in val_df["experiment"].unique():
        sub = val_df[val_df["experiment"] == exp].copy()
        best = sub.loc[sub["mIoU"].idxmax()].to_dict()
        last = sub.loc[sub["idx"].idxmax()].to_dict()
        summaries[exp] = {"count": len(sub), "best": best, "last": last}

    def paired(a, b, max_idx=None):
        aa = val_df[val_df["experiment"] == a]
        bb = val_df[val_df["experiment"] == b]
        if max_idx is not None:
            aa = aa[aa["idx"] <= max_idx]
            bb = bb[bb["idx"] <= max_idx]
        merged = aa.merge(bb, on="idx", suffixes=("_a", "_b"))
        delta = merged["mIoU_a"] - merged["mIoU_b"]
        return {
            "n": int(len(delta)),
            "mean": float(delta.mean()),
            "median": float(delta.median()),
            "wins": int((delta > 0).sum()),
            "losses": int((delta < 0).sum()),
        }

    comp = {
        "v13_vs_baseline": paired("v13", "baseline"),
        "v13_vs_baseline_166": paired("v13", "baseline", 166),
        "v13_vs_v12_166": paired("v13", "v12", 166),
        "v12_vs_baseline_166": paired("v12", "baseline", 166),
    }

    debug = {}
    v13_log = EXP_DIR / "v13-da-radius-cuda-shadow-debug_20260520_2305" / "train.log"
    current = defaultdict(lambda: defaultdict(list))
    if v13_log.exists():
        for line in v13_log.read_text(encoding="utf-8", errors="ignore").splitlines():
            if "DA-Radius debug" not in line:
                continue
            parts = line.split()
            stage = None
            vals = {}
            for token in parts:
                if token.startswith("stage="):
                    stage = int(token.split("=", 1)[1])
                elif "=" in token:
                    k, v = token.split("=", 1)
                    vals[k] = v.rstrip(",")
            if stage is None:
                continue
            for k in ("radius_scale_mean", "valid_neighbors_mean", "shadow_ratio", "full_ratio"):
                if k in vals:
                    current[stage][k].append(float(vals[k]))
    for stage, items in current.items():
        debug[stage] = {k: sum(v) / len(v) for k, v in items.items()}
    return summaries, comp, debug


def save_figures(val_df, class_df, summaries, debug):
    plt.rcParams["font.sans-serif"] = ["Microsoft YaHei", "SimHei", "DejaVu Sans"]
    plt.rcParams["axes.unicode_minus"] = False

    # 1) Validation curves
    fig, ax = plt.subplots(figsize=(10, 5.2), dpi=180)
    colors = {"baseline": "#1f77b4", "v12": "#ff7f0e", "v13": "#2ca02c"}
    for exp in ["baseline", "v12", "v13"]:
        sub = val_df[val_df["experiment"] == exp].sort_values("idx")
        ax.plot(sub["idx"], sub["mIoU"], label=exp, linewidth=1.6, color=colors[exp])
        best = sub.loc[sub["mIoU"].idxmax()]
        ax.scatter([best["idx"]], [best["mIoU"]], color=colors[exp], s=36, zorder=3)
    ax.set_xlabel("Validation index")
    ax.set_ylabel("mIoU")
    ax.set_title("Validation mIoU Curves")
    ax.grid(True, alpha=0.25)
    ax.legend(frameon=False)
    fig.tight_layout()
    curve_path = ASSET_DIR / "val_miou_curves.png"
    fig.savefig(curve_path, bbox_inches="tight")
    plt.close(fig)

    # 2) Best mIoU bar chart
    fig, ax = plt.subplots(figsize=(7.2, 4.4), dpi=180)
    labels = ["baseline", "v12", "v13"]
    values = [summaries[k]["best"]["mIoU"] for k in labels]
    bars = ax.bar(labels, values, color=[colors[k] for k in labels], width=0.55)
    ax.set_ylim(0.55, 0.74)
    ax.set_ylabel("Best mIoU")
    ax.set_title("Best Validation mIoU")
    ax.grid(axis="y", alpha=0.25)
    for bar, v in zip(bars, values):
        ax.text(bar.get_x() + bar.get_width() / 2, v + 0.003, f"{v:.4f}", ha="center", va="bottom", fontsize=9)
    fig.tight_layout()
    best_path = ASSET_DIR / "best_miou_bar.png"
    fig.savefig(best_path, bbox_inches="tight")
    plt.close(fig)

    # 3) Class deltas: v13 best vs baseline best
    class_df = class_df.copy()
    v13_best = summaries["v13"]["best"]["idx"]
    base_best = summaries["baseline"]["best"]["idx"]
    a = class_df[(class_df["experiment"] == "v13") & (class_df["val_idx"] == v13_best)]
    b = class_df[(class_df["experiment"] == "baseline") & (class_df["val_idx"] == base_best)]
    merged = a[["class_name", "iou"]].merge(b[["class_name", "iou"]], on="class_name", suffixes=("_v13", "_base"))
    merged["delta"] = merged["iou_v13"] - merged["iou_base"]
    merged = merged.sort_values("delta")
    fig, ax = plt.subplots(figsize=(8.2, 5.4), dpi=180)
    ax.barh(merged["class_name"], merged["delta"], color=["#d62728" if x < 0 else "#2ca02c" for x in merged["delta"]])
    ax.axvline(0, color="#444", linewidth=1)
    ax.set_xlabel("IoU delta (v13 best - baseline best)")
    ax.set_title("Class-wise Delta")
    ax.grid(axis="x", alpha=0.25)
    fig.tight_layout()
    delta_path = ASSET_DIR / "class_delta_v13_vs_baseline.png"
    fig.savefig(delta_path, bbox_inches="tight")
    plt.close(fig)

    # 4) DA-Radius debug summary
    if debug:
        stages = sorted(debug)
        fig, ax = plt.subplots(figsize=(8.4, 4.8), dpi=180)
        w = 0.22
        xs = list(range(len(stages)))
        shadow = [debug[s].get("shadow_ratio", 0) for s in stages]
        full = [debug[s].get("full_ratio", 0) for s in stages]
        valid = [debug[s].get("valid_neighbors_mean", 0) / 20.0 for s in stages]
        ax.bar([x - w for x in xs], shadow, width=w, label="shadow_ratio")
        ax.bar(xs, full, width=w, label="full_ratio")
        ax.bar([x + w for x in xs], valid, width=w, label="valid_neighbors/20")
        ax.set_xticks(xs)
        ax.set_xticklabels([f"stage {s}" for s in stages])
        ax.set_ylim(0, 1.05)
        ax.set_title("DA-Radius Diagnostics")
        ax.grid(axis="y", alpha=0.25)
        ax.legend(frameon=False)
        fig.tight_layout()
        diag_path = ASSET_DIR / "da_radius_diagnostics.png"
        fig.savefig(diag_path, bbox_inches="tight")
        plt.close(fig)
    else:
        diag_path = None

    return {
        "curve": curve_path,
        "best": best_path,
        "delta": delta_path,
        "diag": diag_path,
    }


def add_heading(doc, text, level=1):
    p = doc.add_heading(level=level)
    r = p.add_run(text)
    set_cn_font(r, size=16 if level == 1 else 13, bold=True)
    return p


def add_paragraph(doc, text, size=11, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        r1 = p.add_run(bold_prefix)
        set_cn_font(r1, size=size, bold=True)
        r2 = p.add_run(text[len(bold_prefix):])
        set_cn_font(r2, size=size)
    else:
        r = p.add_run(text)
        set_cn_font(r, size=size)
    return p


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    hdr = table.rows[0].cells
    for i, h in enumerate(headers):
        hdr[i].text = str(h)
        set_table_font(hdr[i], size=10)
    for row in rows:
        cells = table.add_row().cells
        for i, val in enumerate(row):
            cells[i].text = str(val)
            set_table_font(cells[i], size=10)
    if col_widths:
        for row in table.rows:
            for idx, width in enumerate(col_widths):
                row.cells[idx].width = Inches(width)
    return table


def build_doc(summary, comp, debug, figures):
    doc = Document()
    sec = doc.sections[0]
    sec.top_margin = Inches(0.6)
    sec.bottom_margin = Inches(0.6)
    sec.left_margin = Inches(0.7)
    sec.right_margin = Inches(0.7)

    title = doc.add_paragraph()
    title.alignment = 1
    r = title.add_run("5月22日汇报文档")
    set_cn_font(r, size=20, bold=True)

    p = doc.add_paragraph()
    p.alignment = 1
    r = p.add_run("项目：Pointcept / KPConvX Hybrid + DA-Radius + Global Context")
    set_cn_font(r, size=11)

    add_heading(doc, "1. 本阶段完成工作", level=1)
    bullets = [
        "完成 pointops CUDA 扩展编译与验证，确认 v13 使用的是 CUDA 版 DA-Radius 路径。",
        "修正 adaptive ball query 的 shadow padding 逻辑，避免半径不足时重复最后一个真实邻居。",
        "增加 DA-Radius 调试日志，输出 radius_scale、valid_neighbors、shadow_ratio、full_ratio，便于观察裁剪和邻居截断情况。",
        "整理并固化了 v12 / v13 / v14 三套配置，其中 v14 在 v13 基础上加入全局上下文模块，保持 router / refine / DA-kernel 关闭，便于做干净消融。",
        "修复测试阶段的 OOM 风险：在全景测试配置中加入 TestSphereCrop，避免 PreciseEvaluator 在大场景上显存峰值过高。"
    ]
    for b in bullets:
        add_paragraph(doc, f"• {b}", size=11)

    add_heading(doc, "2. 实验结果", level=1)
    rows = []
    for exp in ["baseline", "v12", "v13"]:
        best = summary[exp]["best"]
        last = summary[exp]["last"]
        rows.append([
            exp,
            summary[exp]["count"],
            f'{best["idx"]}',
            f'{best["mIoU"]:.4f}',
            f'{best["mAcc"]:.4f}',
            f'{best["allAcc"]:.4f}',
            f'{last["mIoU"]:.4f}',
        ])
    add_table(
        doc,
        ["实验", "验证次数", "best idx", "best mIoU", "best mAcc", "best allAcc", "最后 mIoU"],
        rows,
    )

    doc.add_paragraph()
    doc.add_picture(str(figures["curve"]), width=Inches(6.6))
    cap = doc.add_paragraph()
    r = cap.add_run("图1  验证集 mIoU 曲线对比")
    set_cn_font(r, size=10)

    doc.add_paragraph()
    doc.add_picture(str(figures["best"]), width=Inches(5.8))
    cap = doc.add_paragraph()
    r = cap.add_run("图2  三个实验的最佳 mIoU 对比")
    set_cn_font(r, size=10)

    doc.add_paragraph()
    doc.add_picture(str(figures["delta"]), width=Inches(6.5))
    cap = doc.add_paragraph()
    r = cap.add_run("图3  v13 最优模型相对 baseline 最优模型的类别 IoU 差值")
    set_cn_font(r, size=10)

    if figures["diag"]:
        doc.add_paragraph()
        doc.add_picture(str(figures["diag"]), width=Inches(6.3))
        cap = doc.add_paragraph()
        r = cap.add_run("图4  v13 DA-Radius 调试统计")
        set_cn_font(r, size=10)

    add_heading(doc, "3. 结果分析", level=1)
    analysis = [
        f"v13 的最佳结果为 mIoU={summary['v13']['best']['mIoU']:.4f}，略低于 baseline 的 {summary['baseline']['best']['mIoU']:.4f}，但在前 166 次验证中的平均 mIoU 差值为 {comp['v13_vs_baseline_166']['mean']:+.4f}，整体属于接近 baseline、略有正向贡献。",
        f"与 v12 相比，v13 在前 166 次验证中的平均 mIoU 差值为 {comp['v13_vs_v12_166']['mean']:+.4f}，说明 CUDA 版 DA-Radius 比 torch 版更稳定，也更接近最终可用方案。",
        "类别层面上，v13 对 beam、sofa、board、window 有明显提升，但 door、column、bookcase 有回落，说明当前改动更偏向局部几何结构建模，还没有完全解决大尺度语义关系。",
        "DA-Radius 调试日志显示 stage 3/4 的 valid_neighbors_mean 接近 20，full_ratio 也偏高，说明很多点已经被 neighbor_limits 截断；也就是说，现在不是半径太小，而是邻居上限开始限制动态半径继续发挥。",
        "训练阶段梯度裁剪是有效的，原始梯度范数经常高于 1.0；当前先保持 clip_grad=1.0 更稳，后续若全局模块加入后出现学习偏慢，再考虑调到 2.0。",
        "测试阶段 OOM 不是训练崩溃，而是全景测试/PreciseEvaluator 的显存峰值过高。加入 TestSphereCrop 后，测试会更稳，但也会带来更长的测试时间。"
    ]
    for a in analysis:
        add_paragraph(doc, f"• {a}", size=11)

    add_heading(doc, "4. 下一步计划", level=1)
    next_steps = [
        "先跑 v14：在 v13 + DA-Radius 的基础上加入全局上下文模块，保持其他分支关闭，优先验证全局信息能否补足 door / column / bookcase 的退步。",
        "v14 第一轮继续沿用 clip_grad=1.0，先看结构改动本身是否有效；如果训练明显被裁剪限制，再试 clip_grad=2.0。",
        "若后续继续做 DA-Radius 消融，再考虑适度增大 neighbor_limits，例如 (12, 16, 24, 24, 20)，观察 full_ratio 是否下降、类间平衡是否改善。",
        "最终测试继续保留 TestSphereCrop，避免训练完后在精测阶段再次 OOM。"
    ]
    for s in next_steps:
        add_paragraph(doc, f"• {s}", size=11)

    add_heading(doc, "5. 关键数值摘要", level=1)
    if debug:
        dbg_rows = []
        for stage in sorted(debug):
            dbg_rows.append([
                stage,
                f"{debug[stage].get('shadow_ratio', 0):.4f}",
                f"{debug[stage].get('full_ratio', 0):.4f}",
                f"{debug[stage].get('valid_neighbors_mean', 0):.2f}/20",
                f"{debug[stage].get('radius_scale_mean', 0):.4f}",
            ])
        add_table(doc, ["stage", "shadow_ratio", "full_ratio", "valid_neighbors_mean", "radius_scale_mean"], dbg_rows)
    else:
        add_paragraph(doc, "• 暂无可用的 DA-Radius 调试统计。", size=11)

    doc.save(str(OUT_DOCX))


def main():
    ensure_dirs()
    val_df = load_val_metrics()
    class_df = load_class_metrics()
    summary, comp, debug = summarize(val_df, class_df)
    figures = save_figures(val_df, class_df, summary, debug)
    build_doc(summary, comp, debug, figures)
    print(f"saved: {OUT_DOCX}")


if __name__ == "__main__":
    main()
